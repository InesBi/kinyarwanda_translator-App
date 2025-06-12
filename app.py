import streamlit as st
import json
import tempfile
from docx import Document
from fpdf import FPDF
from utils.file_handler import handle_file_upload
from utils.synonym_handler import suggest_synonyms
from services.openai_provider import translate_with_openai
from services.google_provider import translate_with_google
from services.huggingface_provider import translate_with_huggingface

# Clean special characters that FPDF can't encode
def clean_text_for_pdf(text):
    return (
        text.replace("’", "'")
            .replace("“", '"')
            .replace("”", '"')
            .replace("–", "-")
            .replace("—", "-")
            .replace("…", "...")
    )

# Page setup
st.set_page_config(page_title="Umuseke Translator", layout="wide")

# Light styling
st.markdown("""
    <style>
        body {
            font-family: "Segoe UI", sans-serif;
            background-color: #f9f9f9;
        }
        .stTextInput > div > input, .stTextArea textarea {
            font-size: 16px;
        }
        .stDownloadButton button {
            background-color: #25a162;
            color: white;
        }
        section[data-testid="stSidebar"] {
            background: linear-gradient(to bottom, #00ADEF, #20603D, #FAD201);
            color: white;
        }
    </style>
""", unsafe_allow_html=True)

# Title
col1, col2 = st.columns([1, 8])
with col1:
    st.image('https://upload.wikimedia.org/wikipedia/commons/1/17/Flag_of_Rwanda.svg', width=60)
with col2:
    st.markdown("### Umuseke Translator")
    st.markdown(
        "<span style='font-size: 18px; color: gray;'>Translate Kinyarwanda 🇷🇼 to French 🇫🇷 and English 🇬🇧 with context-aware precision.</span>",
        unsafe_allow_html=True
    )

# Load context
with open("config/contexts.json") as f:
    contexts = json.load(f)

# Sidebar controls
provider = st.sidebar.selectbox("Choose Translation Provider", ["OpenAI", "Google", "Hugging Face"])
api_key = st.sidebar.text_input(f"Enter your {provider} API Key", type="password", key="api_key_input")
src_lang = st.sidebar.selectbox("Translate from", ["Kinyarwanda", "English", "French"])
tgt_lang = st.sidebar.selectbox("Translate to", ["English", "French", "Kinyarwanda"])
context = st.sidebar.selectbox("Select Context", list(contexts.keys()))
export_format = st.sidebar.radio("Export format", ["Word (.docx)", "PDF (.pdf)"])

# Input method
input_method = st.radio("Input Method", ["Upload File", "Paste Text"])
input_text = ""
if input_method == "Upload File":
    uploaded_file = st.file_uploader("Upload a text or docx file", type=["txt", "docx"])
    if uploaded_file:
        input_text = handle_file_upload(uploaded_file)
else:
    input_text = st.text_area("Paste your text here")

# Translate
if st.button("Translate"):
    if input_text.strip() == "":
        st.warning("Please provide some input text.")
    else:
        with st.spinner("Translating..."):
            try:
                # Choose translation provider dynamically
                if provider == "OpenAI":
                    translated_text = translate_with_openai(input_text, src_lang, tgt_lang, api_key, context)
                elif provider == "Google":
                    translated_text = translate_with_google(input_text, src_lang, tgt_lang, api_key, context)
                elif provider == "Hugging Face":
                    translated_text = translate_with_huggingface(input_text, src_lang, tgt_lang, api_key, context)
                else:
                    raise ValueError("Unsupported translation provider")

                enriched_text = suggest_synonyms(translated_text, context)
                edited_text = st.text_area("Translated Text (Review & Edit)", value=enriched_text, height=300)

                if export_format == "Word (.docx)":
                    doc = Document()
                    doc.add_heading(f"Translated to {tgt_lang}", level=1)
                    for line in edited_text.split('\n'):
                        doc.add_paragraph(line)
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                        doc.save(tmp.name)
                        with open(tmp.name, "rb") as file:
                            st.download_button(
                                label="📄 Download Word Document",
                                data=file,
                                file_name="translated_output.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )

                elif export_format == "PDF (.pdf)":
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_font("Arial", size=12)
                    for line in edited_text.split('\n'):
                        cleaned_line = clean_text_for_pdf(line)
                        pdf.multi_cell(0, 10, cleaned_line)
                        pdf.ln(2)
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                        pdf.output(tmp.name)
                        with open(tmp.name, "rb") as file:
                            st.download_button(
                                label="📄 Download PDF Document",
                                data=file,
                                file_name="translated_output.pdf",
                                mime="application/pdf"
                            )

                st.success("Translation complete and ready for download!")

            except Exception as e:
                st.error(f"Translation failed: {e}")
